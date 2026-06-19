# -*- coding: utf-8 -*-
"""Shared helpers for thesis Word generation — تنسيق دليل الجامعة."""

from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

META = {
    "title_ar": "إطار رقابة أبوية ذكي: موازنة الأمن الرقمي وعلم نفس الطفل",
    "title_en": "A Stealthy Cyber Parenting Framework: Balancing Digital Security and Child Psychology",
    "student": "[اسم الطالبة]",
    "supervisor": "[اسم المشرف]",
    "university": "[اسم الجامعة]",
    "department": "[قسم نظم المعلومات / علوم الحاسوب]",
    "year": "2025/2026",
}

ARABIC_FONT = "Simplified Arabic"
ENGLISH_FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_run_fonts(run, latin_font: str, cs_font: str, size: int, bold: bool = False):
    run.font.name = latin_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:cs"), cs_font)
    rfonts.set(qn("w:eastAsia"), cs_font)


def font_ar(run, size=14, bold=False):
    _set_run_fonts(run, ARABIC_FONT, ARABIC_FONT, size, bold)


def font_en(run, size=13, bold=False):
    _set_run_fonts(run, ENGLISH_FONT, ENGLISH_FONT, size, bold)


def font(run, size=14, bold=False, color=None):
    """افتراضي عربي — أسود فقط."""
    font_ar(run, size, bold)


def _is_mostly_english(text: str) -> bool:
    letters = re.findall(r"[A-Za-z]", text)
    arabic = re.findall(r"[\u0600-\u06FF]", text)
    return len(letters) > len(arabic)


def para(doc, text, size=14, bold=False, indent=0.75, cite=False, english=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    use_en = english if english is not None else _is_mostly_english(text)
    r = p.add_run(text)
    if use_en:
        font_en(r, 13 if size == 14 else size, bold)
    else:
        font_ar(r, size, bold)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sizes = {1: 16, 2: 15, 3: 14}
    for r in h.runs:
        font_ar(r, sizes.get(level, 14), True)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(item)
        if _is_mostly_english(item):
            font_en(r, 13)
        else:
            font_ar(r, 14)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def table_title(doc, title: str):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(p.add_run(title), 12, True)
    p.paragraph_format.space_after = Pt(4)


def table(doc, headers, rows, title=None):
    if title:
        table_title(doc, title)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                font_ar(r, 12, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                set_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    if _is_mostly_english(str(val)):
                        font_en(r, 12)
                    else:
                        font_ar(r, 12)
    doc.add_paragraph()


def figure(doc, label, caption):
    """عنوان الشكل تحت الشكل — حجم 12."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(p.add_run("[أدرجي الرسم أو لقطة الشاشة هنا]"), 12)
    p.paragraph_format.space_after = Pt(4)
    cap = doc.add_paragraph()
    set_rtl(cap)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(cap.add_run(f"{label}: {caption}"), 12, True)
    cap.paragraph_format.space_after = Pt(10)


def setup_page(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.right_margin = Cm(2.5)
    sec.left_margin = Cm(2.0)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(p.add_run(META["university"]), 18, True)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(p.add_run(META["department"]), 15)
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(p.add_run("مشروع تخرج — درجة البكالوريوس"), 14)
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(p.add_run(META["title_ar"]), 17, True)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_en(p.add_run(META["title_en"]), 13)
    for _ in range(3):
        doc.add_paragraph()
    for line in [
        f"إعداد: {META['student']}",
        f"إشراف: {META['supervisor']}",
        f"العام الجامعي: {META['year']}",
    ]:
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_ar(p.add_run(line), 14)
    doc.add_page_break()


def abstract(doc):
    heading(doc, "الملخص", 1)
    para(
        doc,
        "يهدف هذا البحث إلى تصميم وتنفيذ نظام رقابة أبوية ذكي يوازن بين حماية "
        "الطفل في الفضاء الرقمي والحفاظ على راحته النفسية، عبر دمج آليات المراقبة "
        "داخل واجهة ترفيهية/تعليمية (أكاديمية العباقرة) مع لوحة تحكم لولي الأمر "
        "(MY Rana) وسيرفر سحابي. يعتمد النظام على Kotlin وFlask وربط Gmail حقيقي "
        "ومراقبة Usage Access وAccessibility. أظهرت الاختبارات نجاح الربط والتنبيهات "
        "وسياسات وقت الاستخدام.",
    )
    heading(doc, "Abstract", 2)
    para(
        doc,
        "This research designs and implements a smart parental control framework "
        "balancing digital security and child psychology via a disguised child app "
        "and a parent dashboard, Kotlin/Flask cloud backend, Gmail linking, and "
        "Android usage/accessibility monitoring.",
        english=True,
    )
    doc.add_page_break()


def add_toc(doc, entries):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_ar(p.add_run("فهرس المحتويات"), 18, True)
    p.paragraph_format.space_after = Pt(18)
    for level, title in entries:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.left_indent = Cm(0.4 * level)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        font_ar(p.add_run(title), 13 if level <= 1 else 12, bold=(level == 0))
        p.add_run().add_tab()
        p.add_run(" ")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para(doc, "حدّثي أرقام الصفحات: مراجع ← تحديث الفهرس أو F9", size=11, indent=0)
    doc.add_page_break()
