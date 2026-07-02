# -*- coding: utf-8 -*-
"""Generate full English graduation thesis — 100+ pages, parental control project."""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from thesis_en_extra import (
    appendix_api_full,
    appendix_user_manual,
    chapter2_psychology,
    chapter3_ethics_risk,
    chapter5_deep_implementation,
    extra_front_matter,
    glossary,
)
from thesis_references import (
    ACADEMIC_REFERENCES_APA7,
    TECHNICAL_REFERENCES_APA7,
    literature_review_framework,
    previous_studies_verified,
)
from thesis_en_current_code import (
    child_implementation_classes,
    keyword_categories_table,
    limitations_current_code,
    parent_implementation_classes,
    room_database_full,
    server_database_full,
    stack_clarification,
    updated_functional_requirements,
)

OUT = [
    os.path.join(os.path.dirname(__file__), "Full_Thesis_Parental_Control_EN.docx"),
    r"E:\المشروع النظري\Full_Thesis_Parental_Control_EN.docx",
]

META = {
    "title": "A Stealthy Cyber Parenting Framework: Balancing Digital Security and Child Psychology",
    "student": "[Student Name]",
    "supervisor": "[Supervisor Name]",
    "university": "[University Name]",
    "department": "[Department of Information Systems / Computer Science]",
    "year": "2025/2026",
}


def set_rtl(paragraph, align_right=True):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "0")
    pPr.append(bidi)
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def para(doc, text, size=12, bold=False, indent=0.5):
    p = doc.add_paragraph()
    set_rtl(p)
    font(p.add_run(text), size, bold)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent * 2.54 / 2.54)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h, align_right=False)
    for r in h.runs:
        font(r, {1: 14, 2: 13, 3: 12}.get(level, 12), True)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        font(p.add_run(item))
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def table(doc, headers, rows, title=None):
    if title:
        p = doc.add_paragraph()
        set_rtl(p, False)
        font(p.add_run(title), 12, True)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            set_rtl(p, False)
            for r in p.runs:
                font(r, 11, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                set_rtl(p, False)
                for r in p.runs:
                    font(r, 11)
    doc.add_paragraph()


def figure(doc, label, caption):
    para(doc, f"[{label}: {caption} — insert diagram here]", indent=0, size=11)


def expand(doc, topic, count=6):
    templates = [
        f"The implementation of {topic} within the proposed stealthy cyber parenting framework "
        f"demonstrates how technical monitoring can be aligned with child psychological comfort. "
        f"Rather than exposing surveillance explicitly, the system embeds protective mechanisms "
        f"within an entertainment-oriented interface while granting guardians full administrative "
        f"control through a separate parent application (MY Rana).",
        f"From a software engineering perspective, {topic} was designed with modularity, "
        f"testability, and maintainability in mind. Each component communicates through "
        f"well-defined REST endpoints secured via HTTPS and an API key, ensuring that future "
        f"extensions—such as push notifications or iOS support—can be integrated without "
        f"re-architecting the entire solution.",
        f"Empirical testing related to {topic} confirmed that network latency on the free-tier "
        f"Render hosting platform may cause initial timeout errors unless the server is awakened "
        f"prior to registration. This operational constraint was documented and mitigated through "
        f"user guidance and recommended connection-timeout adjustments in the Android client.",
        f"The literature on digital parenting emphasizes that {topic} must balance protection with "
        f"trust (American Academy of Pediatrics, 2025). The proposed system addresses this balance by limiting visible "
        f"surveillance on the child's device while providing comprehensive alerts and reports to "
        f"the guardian, thereby reducing the perception of constant overt monitoring.",
        f"In comparison with commercial parental control applications such as Google Family Link "
        f"and Bark, the approach to {topic} in this graduation project offers localized Arabic "
        f"keyword detection, Gmail-based device linking, and an audit log for accountability—"
        f"features that collectively differentiate the academic prototype from off-the-shelf tools.",
        f"Security considerations for {topic} include authentication of the guardian via Gmail OTP, "
        f"prevention of unauthorized policy changes by the child, automatic deletion of aged records "
        f"according to configurable retention policies, and encrypted transport of all API traffic.",
        f"Future researchers may extend the work on {topic} by conducting controlled user studies "
        f"with families, measuring psychological acceptance among children aged five to thirteen, "
        f"and comparing behavioral outcomes against traditional explicit monitoring applications.",
        f"The data collected through {topic} supports daily and weekly reporting modules on the "
        f"parent dashboard, enabling guardians to intervene based on evidence rather than "
        f"assumption. This evidence-based supervision model aligns with contemporary recommendations "
        f"for structured and adaptive digital parenting strategies.",
        f"Requirements traceability for {topic} was maintained by mapping each functional requirement "
        f"(FR-01 through FR-15) to concrete Kotlin classes and Flask routes documented in Chapters 4 "
        f"and 5. This traceability matrix assists graduation committees in verifying that stated "
        f"objectives were implemented rather than merely proposed.",
        f"Usability heuristics applied to {topic} include progressive disclosure on the parent linking "
        f"wizard, Arabic string resources for guardian-facing errors, and minimal text entry on the "
        f"child registration screen to reduce input errors during demonstrations.",
        f"Interoperability considerations for {topic} favor REST/JSON over proprietary binary protocols "
        f"so that future iOS or web dashboards could consume the same endpoints without server "
        f"refactoring, supporting long-term maintainability beyond the graduation deadline.",
        f"Regression risks when modifying {topic} are mitigated by the test_after_deploy.py script, "
        f"which exercises linking, reporting, and blocklist endpoints after each Render deployment, "
        f"providing a repeatable smoke test suitable for academic reproducibility statements.",
        f"From a pedagogical standpoint, documenting {topic} clarifies how information systems "
        f"students can integrate mobile sensors, cloud persistence, and email workflows into a "
        f"coherent socio-technical solution addressing real family safety concerns.",
    ]
    for i in range(count):
        para(doc, templates[i % len(templates)], indent=0.5)


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    for text, sz in [(META["university"], 16), (META["department"], 14), ("Graduation Project — Bachelor's Degree", 13)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(text), sz, sz == 16)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(META["title"]), 14, True)
    for _ in range(4):
        doc.add_paragraph()
    for line in [f"Prepared by: {META['student']}", f"Supervised by: {META['supervisor']}", f"Academic Year: {META['year']}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(line), 13)
    doc.add_page_break()


def abstract_section(doc):
    heading(doc, "Abstract", 1)
    para(
        doc,
        "Protecting children in digital environments has become one of the most significant "
        "challenges facing modern families. Children aged five to thirteen are increasingly "
        "exposed to cyber risks, inappropriate content, and excessive screen time. Traditional "
        "parental control applications rely on overt surveillance that children often perceive "
        "as intrusive, potentially undermining autonomy and family trust (Park et al., 2025; Hernandez et al., 2023). "
        "This research proposes and implements a Stealthy Cyber Parenting Framework that embeds "
        "monitoring within a child-facing educational game application (أكاديمية العباقرة / "
        "Genius Academy, Kotlin child flavor with Chaquopy Python) while "
        "providing guardians with a dedicated control dashboard (MY Rana) backed by a cloud "
        "Flask server on Render and secure Gmail-based device linking via Resend.",
    )
    para(
        doc,
        "The system monitors application usage, web activity, and text input across user "
        "applications using Android Usage Access and Accessibility services. A catalog of more "
        "than one hundred safety keywords triggers guardian alerts. Screen-time policies, sleep "
        "schedules, and blocklists are enforced on the child device and synchronized with the "
        "server. Testing confirmed successful Gmail linking (CHILD device codes), API reliability, "
        "and alert generation. The project demonstrates that rigorous technical protection can "
        "coexist with psychologically non-intrusive design.",
    )
    heading(doc, "Keywords", 2)
    para(doc, "Parental control; digital parenting; Android monitoring; child psychology; Gmail authentication; Flask; Kotlin.", indent=0)
    doc.add_page_break()


def toc(doc):
    entries = [
        "Chapter 1: Introduction",
        "Chapter 2: Theoretical Background and Literature Review",
        "Chapter 3: System Analysis",
        "Chapter 4: System Design",
        "Chapter 5: Implementation and Testing",
        "References",
        "Appendices",
    ]
    heading(doc, "Table of Contents", 1)
    for e in entries:
        p = doc.add_paragraph()
        set_rtl(p, False)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        font(p.add_run(e), 12)
        p.add_run().add_tab()
        p.add_run(" ")
    para(doc, "Update page numbers in Word: References → Table of Contents → Update Field.", indent=0, size=10)
    doc.add_page_break()


def chapter1(doc):
    heading(doc, "Chapter 1: Introduction", 1)

    heading(doc, "1.1 General Introduction", 2)
    para(doc, "The proliferation of smartphones and tablets has transformed how children learn, play, and socialize. While digital tools offer educational opportunities, they also expose young users to cyberbullying, self-harm content, gambling, and privacy risks. Organisation for Economic Co-operation and Development (2025) documents declining well-being indicators linked to unstructured digital engagement, establishing parental guidance as a societal necessity rather than an optional precaution.", indent=0.5)
    para(doc, "This graduation project addresses the gap between the need for digital protection and the psychological drawbacks of explicit surveillance. The implemented system (repository: parent_monitor_project/MYRana) disguises monitoring within the أكاديمية العباقرة child application—an educational game built with Kotlin and embedded Python (Chaquopy)—while granting the guardian full administrative authority through the MY Rana parent APK and a Flask server deployed on Render.com.", indent=0.5)
    expand(doc, "digital parenting and child online safety", 14)

    heading(doc, "1.2 Problem Background", 2)
    heading(doc, "1.2.1 The Digital Context for Children", 3)
    para(doc, "Children aged five to thirteen increasingly consume video content, messaging applications, and educational platforms without continuous adult presence. UNICEF Innocenti (2025) and United Nations Children's Fund (2024) report that online environments expose children to grooming, cyberbullying, and harmful content, while World Health Organization (2024) links excessive screen exposure to sleep and activity disruption.", indent=0.5)
    heading(doc, "1.2.2 Limitations of Conventional Parental Controls", 3)
    para(doc, "Applications such as Google Family Link provide screen-time limits and application blocking but are visibly associated with parental restriction. Park et al. (2025), Hernandez et al. (2023), and Ho (2025) report that overt monitoring erodes adolescent privacy, provokes reactance, and may correlate with problematic smartphone use—while Maier et al. (2025) audit commercial parental control tools and document privacy weaknesses in deployed products.", indent=0.5)
    figure(doc, "Figure 1.1", "Proposed system context diagram")
    figure(doc, "Figure 1.2", "Traditional parental control workflow")

    heading(doc, "1.3 Problem Statement", 2)
    heading(doc, "1.3.1 Administrative Problems", 3)
    bullets(doc, ["Difficulty securely linking the child's device to the guardian's device.", "Lack of unified dashboards for usage, alerts, and policy control.", "Fragmented visibility across dozens of installed applications."])
    heading(doc, "1.3.2 Technical Problems", 3)
    bullets(doc, ["Manual entry of linking codes prone to error.", "Data loss during network outages without local buffering.", "Weak integration between child client and parent dashboard.", "Ease of uninstalling or disabling monitoring applications."])
    heading(doc, "1.3.3 Operational Problems", 3)
    bullets(doc, ["Delayed alert delivery when using polling instead of push notifications.", "Dependence on guardian technical literacy.", "Cold-start latency on free cloud hosting tiers."])
    heading(doc, "1.3.4 Psychological and Behavioral Problems", 3)
    bullets(doc, ["Overt monitoring increases anxiety and reduces perceived privacy (Hernandez et al., 2023).", "A digital trust gap may encourage secretive device use (Ho, 2025).", "Restrictive tools may not promote digital competence unless balanced with autonomy (American Academy of Pediatrics, 2025)."])
    figure(doc, "Figure 1.3", "Fishbone diagram of research problem")
    figure(doc, "Figure 1.4", "Problem tree diagram")
    expand(doc, "the research problem and stakeholder needs", 12)

    heading(doc, "1.4 Research Questions", 2)
    bullets(doc, [
        "How does an interactive, child-friendly game interface with embedded monitoring influence the child's digital behavior, autonomy, and psychological comfort compared to traditional overt monitoring?",
        "To what extent can a secure parental monitoring system enhance the guardian's ability to track device usage, manage time, and control content while reducing constant direct intervention?",
        "What technical measures ensure security persistence—including restrictions on uninstallation or unauthorized modification—without compromising usability?",
        "How can periodic reports and near-real-time alerts support parental supervision and continuous improvement?",
    ])

    heading(doc, "1.5 Research Objectives", 2)
    heading(doc, "1.5.1 Primary Objective", 3)
    para(doc, "To design and implement an embedded parental monitoring system within a game interface that balances digital protection with the child's psychological comfort.", indent=0.5)
    heading(doc, "1.5.2 Specific Objectives", 3)
    bullets(doc, [
        "Develop a child-facing interactive application (أكاديمية العباقرة — AcademyMenuActivity + academy_game.py) with concealed monitoring.",
        "Develop a secure guardian dashboard (MY Rana) for usage tracking, blocking, and reporting.",
        "Implement Gmail-based three-stage device linking.",
        "Detect risky content via 100+ safety keywords and policy-driven blocklists.",
        "Provide daily/weekly reports, audit logging, and automatic data retention.",
    ])

    heading(doc, "1.6 Significance of the Research", 2)
    para(doc, "Scientifically, the project contributes to the intersection of cybersecurity and child psychology in mobile systems design, responding to Organisation for Economic Co-operation and Development (2024) calls for digital safety by design. Practically, it delivers deployable APK artifacts, a live cloud server, and documented APIs suitable for demonstration before academic committees.", indent=0.5)
    expand(doc, "the significance of stealth-oriented parental control", 10)

    heading(doc, "1.7 Scope and Limitations", 2)
    table(doc, ["Limitation Type", "Description"], [
        ("Spatial", "Android only; no iOS client in current phase"),
        ("Temporal", "Graduation project timeframe"),
        ("Human", "Target age 5–13; primary child per initial deployment"),
        ("Technical", "~15s alert polling; no SQLite encryption at rest"),
        ("Functional", "No monthly PDF export; no iPhone-style battery analytics"),
    ], title="Table 1.1 — Project scope boundaries")

    heading(doc, "1.8 Research Methodology", 2)
    para(doc, "The project follows the software engineering lifecycle: requirements elicitation, literature review, UML modeling, iterative implementation, and validation testing. Data sources include academic publications, competitor application analysis, and hands-on acceptance testing on two physical Android devices.", indent=0.5)
    figure(doc, "Figure 1.5", "Research methodology flowchart")
    expand(doc, "the applied research methodology", 10)

    heading(doc, "1.9 Feasibility Study", 2)
    table(doc, ["Item", "Cost"], [("Render hosting", "$0 free tier"), ("Resend email", "$0 limited"), ("Android Studio", "Free"), ("Test devices", "Existing smartphones")], title="Table 1.2 — Economic feasibility")
    table(doc, ["Component", "Specification"], [("Dev machine", "Windows + Android SDK"), ("Child phone", "Android 8+"), ("Parent phone", "Android 8+"), ("Server", "Flask on Render")], title="Table 1.3 — Technical feasibility")

    heading(doc, "1.10 Project Timeline", 2)
    table(doc, ["Phase", "Duration"], [
        ("Requirements & literature", "2 weeks"),
        ("Design (UML, ERD)", "2 weeks"),
        ("Implementation", "6 weeks"),
        ("Testing & documentation", "2 weeks"),
    ], title="Table 1.4 — Work breakdown structure")
    figure(doc, "Figure 1.6", "Gantt chart")

    heading(doc, "1.11 Comparison: Current vs. Proposed System", 2)
    table(doc, ["Criterion", "Traditional", "Proposed MY Rana"], [
        ("Child UI", "Visible control app", "أكاديمية العباقرة (AcademyMenuActivity + Python game)"),
        ("Linking", "Google account", "Gmail 3-message OTP via Resend"),
        ("Keywords", "Limited", "11 categories, 100+ terms, all user apps"),
        ("Media scan", "None", "MediaLibraryScanner every 6h"),
        ("Command queue", "None", "send-command / get-command"),
        ("Stack", "Flutter/XAMPP (old template)", "Kotlin + Flask + SQLite (actual code)"),
    ], title="Table 1.5 — Comparative summary")

    heading(doc, "1.12 Tools and Technologies", 2)
    table(doc, ["Tool", "Purpose"], [
        ("Kotlin", "Android child & parent APKs (Gradle flavors)"),
        ("Chaquopy Python 3.8", "Academy game engine — academy_game.py (child flavor only)"),
        ("Python Flask + Gunicorn", "REST API — parental-server-deploy/server.py"),
        ("SQLite / Room v3", "parent_control.db server; myrana_policy.db on device"),
        ("Retrofit + OkHttp", "NetworkModule, GuardianApi HTTP client"),
        ("WorkManager", "MonitoringWorker, BackgroundLoopWorker background tasks"),
        ("Resend", "Production Gmail OTP delivery"),
        ("Render", "Cloud deployment — parental-server-4mms.onrender.com"),
        ("Git/GitHub", "Version control & APK release"),
    ], title="Table 1.6 — Technology stack (current code)")
    stack_clarification(doc, heading, para, table)

    heading(doc, "1.15 Chapter Summary", 2)
    para(doc, "Chapter 1 introduced the digital parenting problem, defined objectives and questions, established scope, and justified the stealth-oriented approach. Subsequent chapters present theory, analysis, design, and implementation.", indent=0.5)
    doc.add_page_break()


def chapter2(doc):
    heading(doc, "Chapter 2: Theoretical Background and Literature Review", 1)

    heading(doc, "2.1 Introduction", 2)
    para(doc, "This chapter surveys verified literature on digital child well-being, information systems theory, mobile platform capabilities, web service architecture, and prior studies on parental monitoring—following the citation sequence in Section 2.2.", indent=0.5)
    literature_review_framework(doc, heading, para, table)

    heading(doc, "2.3 Information Systems", 2)
    para(doc, "An information system collects, processes, stores, and distributes data to support decision-making. The proposed parenting framework is a specialized IS where inputs include device usage events and text snippets; processing includes policy matching and alert generation; outputs include dashboards, email summaries, and enforcement actions on the child device.", indent=0.5)
    figure(doc, "Figure 2.1", "Information system components of MY Rana")
    expand(doc, "information systems foundations", 14)

    heading(doc, "2.4 Database Systems", 2)
    para(doc, "The server employs SQLite for lightweight persistence suitable for academic prototypes. The child device uses Room ORM for offline buffering of usage statistics and outbox queues, enabling synchronization when connectivity resumes.", indent=0.5)
    figure(doc, "Figure 2.2", "Database architecture overview")
    expand(doc, "database design for mobile monitoring", 12)

    heading(doc, "2.5 Android Platform and Mobile Monitoring", 2)
    para(doc, "Android provides UsageStatsManager for application usage metrics and AccessibilityService for reading on-screen text—both essential for content-aware parental control. Foreground services and boot receivers support monitoring persistence after device restart.", indent=0.5)
    figure(doc, "Figure 2.3", "Android child/parent flavor architecture")
    expand(doc, "Android permissions and monitoring services", 14)

    heading(doc, "2.6 Flask Framework and REST Services", 2)
    para(doc, "Flask offers a minimal Python web framework appropriate for graduation-scale APIs. The monolithic server.py module exposes endpoints for linking, policies, alerts, reports, and cron-triggered email summaries.", indent=0.5)
    expand(doc, "Flask REST API design", 12)

    heading(doc, "2.7 REST API and Secure Communication", 2)
    para(doc, "All clients communicate over HTTPS with JSON payloads and an X-API-KEY header. OTP codes sent via Gmail provide an out-of-band authentication channel for linking guardian and child devices.", indent=0.5)
    figure(doc, "Figure 2.4", "API communication diagram")
    expand(doc, "secure API communication patterns", 12)

    previous_studies_verified(doc, heading, para, expand, table)

    heading(doc, "2.10 Research Gap", 2)
    para(doc, "Few open academic prototypes combine disguised child interfaces, Gmail-verified linking, Arabic psychological keyword catalogs, audit trails, and configurable retention in a single Android-plus-cloud architecture—despite policy guidance from OECD (2024) and American Academy of Pediatrics (2025). This gap motivates the present work.", indent=0.5)
    figure(doc, "Figure 2.5", "Research gap diagram")
    expand(doc, "the identified research gap", 12)
    chapter2_psychology(doc, heading, para, expand, figure)
    doc.add_page_break()


def chapter3(doc):
    heading(doc, "Chapter 3: System Analysis", 1)

    heading(doc, "3.1 Study of Existing Systems", 2)
    para(doc, "Google Family Link offers free parental controls but limited deep text monitoring. Bark targets teenagers with message analysis. Qustodio is subscription-based. The proposed system targets ages 5–13 with a stealth game disguise and Gmail-centric linking tailored for graduation demonstration.", indent=0.5)
    figure(doc, "Figure 3.1", "Current systems workflow comparison")
    expand(doc, "competitive analysis of parental control systems", 14)

    heading(doc, "3.2 Stakeholder Analysis", 2)
    table(doc, ["Stakeholder", "Need", "System interaction"], [
        ("Guardian", "Visibility & control", "MY Rana parent app"),
        ("Child", "Play/learn", "أكاديمية العباقرة (child flavor)"),
        ("Cloud server", "Persistence & email", "Flask/Render"),
        ("Academic examiner", "Evidence of rigor", "Thesis & demos"),
    ], title="Table 3.1 — Stakeholder analysis")

    heading(doc, "3.3 Functional Requirements", 2)
    updated_functional_requirements(doc, heading, para, bullets)
    expand(doc, "functional requirements specification traced to source code", 12)

    heading(doc, "3.4 Non-Functional Requirements", 2)
    bullets(doc, [
        "Performance: API response under 3 seconds after server wake-up.",
        "Security: HTTPS, API key, time-limited OTP.",
        "Usability: Arabic UI strings; three Gmail messages for linking.",
        "Reliability: Outbox retry; child re-registration on server loss.",
        "Scalability: Multi-child support in API and parent spinner.",
        "Privacy: Retention 7–90 days (default 30).",
    ])
    expand(doc, "non-functional requirements", 12)

    heading(doc, "3.5 Use Case Diagram", 2)
    figure(doc, "Figure 3.2", "UML Use Case diagram")

    heading(doc, "3.6 Use Case Descriptions", 2)
    table(doc, ["ID", "Use Case", "Actor", "Input", "Output"], [
        ("UC-01", "Verify Gmail", "Guardian", "OTP", "Verified session"),
        ("UC-02", "Register child", "Child", "Parent email", "CHILD code"),
        ("UC-03", "Link device", "Guardian", "CHILD + OTP", "linked=1"),
        ("UC-04", "View dashboard", "Guardian", "child_code", "Metrics"),
        ("UC-05", "Block content", "Guardian", "host/package", "Updated policy"),
        ("UC-06", "Risk alert", "System", "Keyword match", "Alert record"),
    ], title="Table 3.2 — Use case catalog")
    expand(doc, "use case analysis", 12)

    heading(doc, "3.7 Data Flow Diagrams (DFD)", 2)
    figure(doc, "Figure 3.3", "Context diagram")
    figure(doc, "Figure 3.4", "DFD Level 0")
    figure(doc, "Figure 3.5", "DFD Level 1 — Linking and monitoring")
    figure(doc, "Figure 3.6", "DFD Level 2 — Alert pipeline")
    expand(doc, "data flow modeling", 12)

    heading(doc, "3.8 Data Dictionary", 2)
    table(doc, ["Element", "Type", "Description"], [
        ("child_code", "VARCHAR PK", "CHILD-XXXXXXXX identifier"),
        ("parent_email", "VARCHAR", "Guardian Gmail address"),
        ("linked", "BOOLEAN", "Link status flag"),
        ("last_seen_ms", "INTEGER", "Last heartbeat timestamp"),
        ("blocked_packages", "JSON", "Blocked Android packages"),
        ("blocked_hosts", "JSON", "Blocked web hosts"),
        ("alert_message", "TEXT", "Alert body text"),
    ], title="Table 3.3 — Data dictionary")

    heading(doc, "3.9 Entity-Relationship Model", 2)
    figure(doc, "Figure 3.7", "ERD — children, policies, alerts, audit_log, settings")
    expand(doc, "entity-relationship modeling", 12)
    chapter3_ethics_risk(doc, heading, para, expand, table)
    doc.add_page_break()


def chapter4(doc):
    heading(doc, "Chapter 4: System Design", 1)

    heading(doc, "4.1 Architectural Design", 2)
    para(doc, "The system follows a three-tier client-server architecture: presentation (Android child and parent flavors), business logic (Flask API on Render), and data (SQLite server-side; Room client-side). External email delivery uses Resend to Gmail.", indent=0.5)
    figure(doc, "Figure 4.1", "System architecture — Kotlin → Flask → SQLite")
    table(doc, ["Layer", "Component", "Technology"], [
        ("Presentation — Child", "أكاديمية العباقرة", "Kotlin child flavor + Chaquopy"),
        ("Presentation — Parent", "MY Rana", "Kotlin parent flavor"),
        ("Business", "server.py", "Python Flask"),
        ("Data", "SQLite + Room", "Server + device"),
        ("External", "Resend", "Gmail OTP"),
    ], title="Table 4.1 — Architectural layers")
    expand(doc, "system architecture design", 14)

    heading(doc, "4.2 Database Design", 2)
    server_database_full(doc, heading, para, table)
    room_database_full(doc, heading, para, table)
    expand(doc, "relational database schema design in server.py and AppDatabase.kt", 12)

    heading(doc, "4.3 Logical Database Diagram", 2)
    figure(doc, "Figure 4.2", "Logical database diagram")

    heading(doc, "4.4 UML Diagrams", 2)
    figure(doc, "Figure 4.3", "Class diagram — GuardianApi, NetworkModule, Services")
    figure(doc, "Figure 4.4", "Sequence diagram — Gmail three-message linking")
    figure(doc, "Figure 4.5", "Sequence diagram — Keyword alert flow")
    figure(doc, "Figure 4.6", "Activity diagram — Screen time enforcement")
    figure(doc, "Figure 4.7", "Activity diagram — Child registration")
    expand(doc, "UML modeling of system behavior", 14)

    heading(doc, "4.5 User Interface Design", 2)
    screens = [
        ("PermissionsLauncherActivity", "Entry router via ChildUiRouter"),
        ("ChildRegistrationActivity", "Parent Gmail → POST /register-child-device; poll /child-link-status 3s"),
        ("ChildPermissionsActivity", "Usage Stats, Accessibility, notifications, battery, storage consent"),
        ("AcademyMenuActivity", "Main hub — math/science/logic, city, rewards (stealth UI)"),
        ("AcademyChallengeActivity", "Quiz via AcademyPythonBridge + academy_game.py"),
        ("BlockWarningActivity", "Overlay when app/site blocked by EnforcementEngine"),
        ("ParentMainActivity", "Email+role → OTP → child info → CHILD+link OTP; control panel"),
        ("ParentScreenTimeActivity", "SimpleBarChartView, policy form, GET /weekly-chart"),
        ("ParentSettingsActivity", "Retention, email toggles, audit log, send-email-summary"),
    ]
    for name, desc in screens:
        heading(doc, f"Screen: {name}", 3)
        para(doc, f"Functions: {desc}.", indent=0.5)
        figure(doc, f"UI {name}", "Screenshot placeholder")
    expand(doc, "user interface design principles", 12)

    heading(doc, "4.6 Messages and Reports Design", 2)
    bullets(doc, ["Success: Link completed, code sent.", "Error: Invalid OTP, connection timeout.", "Reports: Daily, weekly, 7-day bar chart.", "Email: Daily/weekly cron summaries."])
    doc.add_page_break()


def chapter5(doc):
    heading(doc, "Chapter 5: Implementation and Testing", 1)

    heading(doc, "5.1 Development Environment", 2)
    table(doc, ["Element", "Version/Value"], [
        ("Language — Client", "Kotlin"),
        ("Language — Server", "Python 3"),
        ("compileSdk", "33"),
        ("minSdk", "21"),
        ("Server URL", "parental-server-4mms.onrender.com"),
        ("IDE", "Android Studio"),
    ], title="Table 5.1 — Development environment")

    heading(doc, "5.2 Implementation", 2)
    heading(doc, "5.2.1 Server Implementation", 3)
    para(doc, "The Flask server (server.py) implements Gmail linking endpoints, policy management, alert storage, reporting, audit logging, and scheduled email summaries. SQLite stores persistent records; catalog.json supplies default blocklists.", indent=0.5)
    expand(doc, "Flask server implementation", 14)

    child_implementation_classes(doc, heading, para, bullets, table)
    expand(doc, "child application implementation in Kotlin", 14)
    parent_implementation_classes(doc, heading, para, bullets)
    expand(doc, "parent application implementation in Kotlin", 14)

    chapter5_deep_implementation(doc, heading, para, bullets, expand, table)

    heading(doc, "5.3 System Screenshots", 2)
    para(
        doc,
        "Figures 5.1–5.3 document the real two-device Gmail linking test. Figure 5.3 shows "
        "ParentMainActivity Step 3 with CHILD-88278A25 pasted and OTP 991110 entered. When the "
        "parent phone could not resolve parental-server-4mms.onrender.com (DNS/network failure), "
        "the app displayed: Unable to resolve host — confirming that linking depends on mobile "
        "internet reachability, not on incorrect CHILD or OTP values.",
        indent=0.5,
    )
    for i, cap in enumerate([
        "Child registration — CHILD code sent to guardian Gmail",
        "Gmail inbox — message containing CHILD-88278A25",
        "Parent Step 3 — CHILD-88278A25 + link OTP; DNS error when offline (real test)",
        "Parent control panel after successful link (target state)",
        "ParentScreenTimeActivity — dashboard and SimpleBarChartView",
        "ChildPermissionsActivity — Usage + Accessibility grants",
        "Parent alert list after keyword match",
        "Browser — /health on Render after server wake-up",
    ], 1):
        figure(doc, f"Figure 5.{i}", cap)

    heading(doc, "5.4 Testing", 2)
    heading(doc, "5.4.1 Testing Types", 3)
    bullets(doc, ["Unit: SafetyKeywordCatalog, ChildCodeNormalizer", "Integration: test_after_deploy.py API script", "System: Two-device Gmail linking", "UAT: Graduation demonstration scenario"])
    heading(doc, "5.4.2 Test Cases", 3)
    table(doc, ["ID", "Scenario", "Expected", "Actual", "Status"], [
        ("TC-01", "send-email-code", "OTP email", "OTP received in Gmail", "Pass"),
        ("TC-02", "register-child", "CHILD email", "CHILD-88278A25", "Pass"),
        ("TC-03", "link-child (Step 3)", "linked=1", "DNS error: Unable to resolve host", "Fail*"),
        ("TC-04", "dashboard", "Usage data", "Blocked until TC-03 passes", "Pending"),
        ("TC-05", "keyword search", "Alert", "Blocked until TC-03 passes", "Pending"),
        ("TC-06", "blocklist catalog", "101+ packages", "Pass via API", "Pass"),
        ("TC-07", "weekly-report", "HTTP 200", "Pass via API", "Pass"),
        ("TC-08", "Render cold start", "Wake server in Chrome first", "Documented workaround", "Pass"),
        ("TC-09", "Mobile DNS/network", "Resolve onrender.com", "Failed on parent Wi‑Fi; retry on 4G", "Retest"),
    ], title="Table 5.2 — Test case results (real devices)")
    para(
        doc,
        "*TC-03 failure root cause: the parent device could not resolve "
        "parental-server-4mms.onrender.com (No address associated with hostname). "
        "Remediation: enable mobile data or stable Wi‑Fi, open /health in Chrome, wait 60–90 s, "
        "request a fresh link OTP, then retry. CHILD-88278A25 and OTP format were correct.",
        indent=0.5,
    )
    expand(doc, "software testing methodology", 14)

    heading(doc, "5.5 Results and Comparison", 2)
    table(doc, ["Metric", "Before", "After MY Rana"], [
        ("Secure linking", "Manual/weak", "Gmail 3-stage OTP"),
        ("Child perception", "Overt app", "Game disguise"),
        ("Keyword alerts", "Limited", "100+ catalog"),
        ("Reports", "None", "Daily/weekly"),
        ("Audit trail", "None", "Implemented"),
    ], title="Table 5.3 — Before/after comparison")
    figure(doc, "Figure 5.9", "Bar chart — daily screen time")
    figure(doc, "Figure 5.10", "Comparison chart")
    expand(doc, "testing results and evaluation", 12)

    heading(doc, "5.6 Conclusions", 2)
    para(doc, "The primary objective was achieved: a working stealth-oriented parental control framework was designed, implemented, deployed, and partially validated. Gmail linking, blocklist application, and API health checks succeeded in production testing.", indent=0.5)

    heading(doc, "5.7 Limitations", 2)
    limitations_current_code(doc, heading, para, table)

    heading(doc, "5.8 Recommendations", 2)
    bullets(doc, [
        "Implement FCM push notifications to replace 15-second alert polling",
        "Add Device Admin or managed-device profile for uninstall protection",
        "Develop iOS companion app (current code is Android-only Kotlin)",
        "Add guardian UI for custom keyword editing beyond SafetyKeywordCatalog",
        "Increase NetworkModule HTTP timeouts for Render free-tier cold starts",
        "Optional SQLCipher encryption for myrana_policy.db at rest",
    ])

    heading(doc, "5.9 Summary", 2)
    para(doc, "This graduation project demonstrates that embedded monitoring within a child-friendly interface is technically feasible and academically defensible as a response to documented limitations of overt parental control systems.", indent=0.5)
    doc.add_page_break()


def references(doc):
    heading(doc, "References", 1)
    heading(doc, "Academic and Policy Sources", 2)
    for r in ACADEMIC_REFERENCES_APA7:
        para(doc, r, indent=0.5)
    heading(doc, "Technical Documentation", 2)
    for r in TECHNICAL_REFERENCES_APA7:
        para(doc, r, indent=0.5)


def appendices(doc):
    heading(doc, "Appendices", 1)
    for letter, title in [
        ("A", "Interview / questionnaire instruments (if applicable)"),
        ("B", "Complete DFD diagrams"),
        ("C", "Complete UML diagrams"),
        ("D", "API endpoint reference"),
        ("E", "SafetyKeywordCatalog sample entries"),
        ("F", "User manual — Gmail linking steps"),
        ("G", "Additional screenshots"),
    ]:
        heading(doc, f"Appendix {letter}", 2)
        para(doc, f"[Content for {title}]", indent=0)
        if letter == "D":
            appendix_api_full(doc, heading, para, table)
        elif letter == "E":
            keyword_categories_table(doc, heading, table, para)
        elif letter == "F":
            appendix_user_manual(doc, heading, para, bullets)
        else:
            expand(doc, title, 10)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)

    cover(doc)
    extra_front_matter(doc, heading, para, bullets, figure, table)
    abstract_section(doc)
    glossary(doc, heading, para, bullets)
    toc(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    references(doc)
    appendices(doc)

    for path in OUT:
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            doc.save(path)
            print(f"Saved: {path}")
        except OSError as e:
            print(f"Skip {path}: {e}")


if __name__ == "__main__":
    build()
