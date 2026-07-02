# -*- coding: utf-8 -*-
"""Generate Chapter 5 (Implementation) — standalone Word document."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_PATHS = [
    os.path.join(os.path.dirname(__file__), "الفصل_الخامس_التنفيذ.docx"),
    r"E:\المشروع النظري\الفصل_الخامس_التنفيذ.docx",
]


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def font(run, size=14, bold=False):
    run.font.name = "Traditional Arabic"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Traditional Arabic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Traditional Arabic")
    run._element.rPr.rFonts.set(qn("w:cs"), "Traditional Arabic")


def para(doc, text, size=14, bold=False, indent=0.75):
    p = doc.add_paragraph()
    set_rtl(p)
    r = p.add_run(text)
    font(r, size, bold)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for r in h.runs:
        font(r, {1: 16, 2: 15, 3: 14}.get(level, 14), True)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        font(p.add_run(item))
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            set_rtl(p)
            for r in p.runs:
                font(r, 12, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                set_rtl(p)
                for r in p.runs:
                    font(r, 12)
    doc.add_paragraph()


def build_chapter5(doc):
    heading(doc, "الفصل الخامس: التنفيذ والبرمجة", 1)

    para(
        doc,
        "يتناول هذا الفصل الجوانب العملية لتنفيذ نظام الرقابة الأبوية المقترح، "
        "المبني على فكرة دمج المراقبة داخل واجهة ترفيهية للطفل (أكاديمية العباقرة) "
        "مع لوحة تحكم مستقلة لولي الأمر (MY Rana). يصف الفصل بيئة التطوير، "
        "هيكل المشروع، تنفيذ تطبيقي Android، السيرفر السحابي، آلية الربط عبر Gmail، "
        "ومحركات المراقبة والتنبيه. جميع ما ورد هنا يعكس التنفيذ الفعلي في المشروع "
        "وليس مجرد تصور نظري.",
    )

    # 5.1
    heading(doc, "5.1 مقدمة التنفيذ", 2)
    para(
        doc,
        "اعتمد التنفيذ بنية Client-Server: طبقة عرض على جهازين Android (نكهتان "
        "من نفس المستودع)، وطبقة خدمات على سيرفر Flask مستضاف على Render، "
        "مع خدمة بريد Resend لإرسال رسائل Gmail. لغة العميل Kotlin، ولغة "
        "السيرفر Python 3. الاتصال عبر HTTPS مع رأس X-API-KEY لكل طلب.",
        indent=0.75,
    )
    table(
        doc,
        ["المكوّن", "التقنية المنفّذة"],
        [
            ("تطبيق الطفل", "Kotlin — flavor: child — أكاديمية العباقرة"),
            ("تطبيق ولي الأمر", "Kotlin — flavor: parent — MY Rana"),
            ("السيرفر", "Python Flask — parental-server-4mms.onrender.com"),
            ("قاعدة البيانات", "SQLite على السيرفر + Room محلياً على الطفل"),
            ("البريد", "Resend API → Gmail (ربط ثلاثي المراحل)"),
            ("التحكم بالإصدارات", "Git / GitHub"),
            ("IDE", "Android Studio + Cursor/VS Code"),
        ],
    )

    # 5.2
    heading(doc, "5.2 هيكل مشروع Android (MYRana)", 2)
    para(
        doc,
        "يستخدم المشروع Gradle Product Flavors لبناء تطبيقين منفصلين دون تكرار "
        "الكود المشترك. كل تطبيق له applicationId مستقل: com.example.myrana.child "
        "و com.example.myrana.parent، مما يمنع الطفل من الوصول لواجهة ولي الأمر "
        "على نفس الجهاز بسهولة.",
        indent=0.75,
    )
    bullets(
        doc,
        [
            "src/main/: الكود المشترك (شبكة، مراقبة، قاعدة محلية، أكاديمية).",
            "src/parent/: واجهات ولي الأمر فقط (ParentMainActivity، الإعدادات، التقارير).",
            "src/child/: موارد وإعدادات نكهة الطفل إن وُجدت.",
            "BuildConfig: SERVER_ROOT_URL، SERVER_BASE_URL، API_KEY مضمّنة في APK.",
        ],
    )

    # 5.3 Child app
    heading(doc, "5.3 تنفيذ تطبيق الطفل (الواجهة المموّهة)", 2)
    heading(doc, "5.3.1 التسجيل والربط", 3)
    para(
        doc,
        "يبدأ الطفل من ChildRegistrationActivity: إدخال بريد ولي الأمر (Gmail)، "
        "ثم POST /register-child-device. عند نجاح الإرسال email_sent=true يُخفى "
        "كود CHILD من الشاشة ويُرسل إلى Gmail فقط — بما يتوافق مع فكرة عدم "
        "إزعاج الطفل بنسخ الأكواد. يستمر التطبيق في polling على /child-link-status "
        "كل 3 ثوانٍ حتى linked=true، ثم ينتقل إلى ChildPermissionsActivity.",
        indent=0.75,
    )
    heading(doc, "5.3.2 واجهة الأكاديمية (اللعبة)", 3)
    para(
        doc,
        "بعد منح الصلاحيات يصل الطفل إلى AcademyMenuActivity وأنشطة الألعاب "
        "التعليمية. هذه الواجهة تحقق الجانب النفسي للمشروع: الطفل يرى تطبيقاً "
        "ترفيهياً/تعليمياً بينما تعمل خدمات المراقبة في الخلفية دون شاشات "
        "تذكير بالمراقبة المستمرة.",
        indent=0.75,
    )
    heading(doc, "5.3.3 الصلاحيات", 3)
    table(
        doc,
        ["الصلاحية", "الغرض", "الإلزام"],
        [
            ("Usage Access", "إحصائيات وقت التطبيقات", "إلزامي"),
            ("Accessibility Service", "قراءة النصوص والروابط", "إلزامي"),
            ("Notifications", "تنبيهات محلية للحدود", "مستحسن"),
            ("Battery optimization", "استمرار الخدمة بالخلفية", "مستحسن"),
            ("READ_MEDIA_*", "فحص ملفات الوسائط", "مستحسن"),
        ],
    )

    # 5.4 Monitoring
    heading(doc, "5.4 محرك المراقبة على جهاز الطفل", 2)
    heading(doc, "5.4.1 ParentSyncService (خدمة أمامية)", 3)
    para(
        doc,
        "ParentSyncService تعمل كـ Foreground Service لضمان بقاء المراقبة نشطة. "
        "ترفع heartbeat إلى /child-heartbeat لتحديث last_seen، وتجمع بيانات "
        "الاستخدام عبر ScreenTimeSyncHelper و UsageUploadHelper، وتُطلق "
        "MediaLibraryScanner و BackgroundMonitoring دورياً.",
        indent=0.75,
    )
    heading(doc, "5.4.2 ContentFilterAccessibilityService", 3)
    para(
        doc,
        "تفحص نصوص جميع تطبيقات المستخدم (عبر MonitoredAppRegistry) مقابل "
        "SafetyKeywordCatalog (أكثر من 100 كلمة في فئات: إيذاء النفس، تنمر، "
        "مخدرات، قمار، إلخ) وكذلك video_keywords من سياسة السيرفر. عند "
        "اكتشاف كلمة خطرة أو موقع محظور تُرسل POST /add-alert وتُعرض "
        "BlockWarningActivity على جهاز الطفل.",
        indent=0.75,
    )
    heading(doc, "5.4.3 EnforcementEngine", 3)
    para(
        doc,
        "يراقب التطبيق الأمامي (foreground) ويطبّق الحظر من PolicyFilterCache "
        "(blocked_packages، blocked_hosts). يتكامل مع AppUsageAlertHelper "
        "لإرسال تنبيه عند تغيير التطبيق (مع cooldown 5 دقائق لكل تطبيق).",
        indent=0.75,
    )
    heading(doc, "5.4.4 MediaLibraryScanner", 3)
    para(
        doc,
        "يفحص MediaStore (صور، فيديو، صوت) بحثاً عن أسماء ملفات أو بيانات "
        "وصفية تحتوي كلمات خطرة، ويرسل تنبيهات للسيرفر عند المطابقة.",
        indent=0.75,
    )
    heading(doc, "5.4.5 ScreenTimeEnforcer ووقت النوم", 3)
    para(
        doc,
        "يقرأ سياسة الحدود (warn_minutes، block_minutes، sleep_start/end) من "
        "السيرفر أو التخزين المحلي. يُظهر تحذيرات متدرجة (أصفر/أحمر) ويطبّق "
        "ScreenTimeSleepHelper لرصد الاستخدام بعد وقت النوم وإرسال تنبيهات.",
        indent=0.75,
    )
    heading(doc, "5.4.6 التخزين المحلي والمزامنة", 3)
    para(
        doc,
        "يستخدم التطبيق Room (AppDatabase) لتخزين أحداث وقت الشاشة واستخدام "
        "التطبيقات. OutboxRepository يحفظ الرفوعات الفاشلة ويعيد إرسالها عند "
        "عودة الشبكة — بما يتوافق مع متطلبات العمل دون اتصال دائم.",
        indent=0.75,
    )

    # 5.5 Parent app
    heading(doc, "5.5 تنفيذ تطبيق ولي الأمر", 2)
    para(
        doc,
        "ParentMainActivity يدير تدفقاً من أربع خطوات: (1) بريد وصفة ولي الأمر، "
        "(2) تحقق OTP من Gmail، (3) اسم الطفل وربط الجهاز بكود CHILD ورمز الربط، "
        "(4) لوحة التحكم. GuardianApi يغلف كل استدعاءات REST. بعد الربط يعمل "
        "polling للتنبيهات كل 15 ثانية في onResume.",
        indent=0.75,
    )
    bullets(
        doc,
        [
            "ParentScreenTimeActivity: لوحة مؤشرات يومية + رسوم أسبوعية.",
            "ParentSettingsActivity: حدود الوقت، وقت النوم، retention، audit log.",
            "أوامر فورية: حظر موقع/تطبيق، تجميد، جدولة، قائمة حظر كاملة.",
            "دعم عدة أطفال: /list-children + Spinner في الواجهة.",
        ],
    )

    # 5.6 Server
    heading(doc, "5.6 تنفيذ السيرفر (server.py)", 2)
    para(
        doc,
        "ملف server.py monolithic Flask يحتوي على: مسارات الربط (send-email-code، "
        "register-child-device، send-link-code، link-child)، مسارات السياسة "
        "والحظر، التقارير (daily-report، weekly-report، weekly-chart)، التنبيهات "
        "(add-alert، alerts)، إعدادات ولي الأمر (guardian-settings)، audit-log، "
        "وتنظيف البيانات القديمة _cleanup_old_data حسب retention_days.",
        indent=0.75,
    )
    heading(doc, "5.6.1 قاعدة البيانات", 3)
    table(
        doc,
        ["الجدول", "الوظيفة"],
        [
            ("children", "بيانات الطفل والربط و last_seen"),
            ("device_policies", "حزم ومواقع وكلمات محظورة"),
            ("alerts", "تنبيهات من جهاز الطفل"),
            ("audit_log", "سجل العمليات الحساسة"),
            ("guardian_settings", "إعدادات ولي الأمر"),
            ("email_otps", "رموز Gmail المؤقتة"),
        ],
    )
    heading(doc, "5.6.2 قائمة الحظر الافتراضية", 3)
    para(
        doc,
        "ملف blocklists/catalog.json يحتوي 101+ حزمة تطبيق ومواقع وكلمات فيديو. "
        "يُطبَّق على الجهاز عبر /apply-default-blocklist بعد الربط.",
        indent=0.75,
    )

    # 5.7 Gmail
    heading(doc, "5.7 تنفيذ الربط عبر Gmail", 2)
    para(
        doc,
        "يحقق النظام email_real_linking=true و dev_fallback_enabled=false على "
        "الإنتاج. تسلسل الرسائل الثلاث:",
        indent=0.75,
    )
    bullets(
        doc,
        [
            "الرسالة 1: OTP تحقق بريد ولي الأمر (/send-email-code).",
            "الرسالة 2: كود CHILD-XXXXXXXX (/register-child-device).",
            "الرسالة 3: OTP ربط الجهاز (/send-link-code ثم /link-child).",
        ],
    )
    para(
        doc,
        "عند نجاح إرسال CHILD بالبريد يُستبعد child_code من استجابة JSON "
        "لمنع النسخ من واجهة برمجة التطبيقات — ولي الأمر يلصق من Gmail فقط.",
        indent=0.75,
    )

    # 5.8 Security
    heading(doc, "5.8 الأمان والخصوصية في التنفيذ", 2)
    bullets(
        doc,
        [
            "HTTPS إلزامي بين التطبيقات و Render.",
            "X-API-KEY على كل طلب (@app.before_request).",
            "OTP محدود زمنياً في قاعدة البيانات.",
            "فصل applicationId بين الطفل والأم.",
            "Audit Log لعمليات الحظر وتغيير الإعدادات.",
            "حذف تلقائي للسجلات بعد retention_days (افتراضي 30 يوماً).",
        ],
    )

    # 5.9 Deploy
    heading(doc, "5.9 النشر والتوزيع", 2)
    para(
        doc,
        "السيرفر منشور على Render.com ومتصل بـ Resend لـ Gmail. ملفات APK "
        "debug مبنية عبر assembleChildDebug و assembleParentDebug ومتوفرة في "
        "GitHub releases: app-child-debug.apk و app-parent-debug.apk.",
        indent=0.75,
    )
    table(
        doc,
        ["العنصر", "القيمة"],
        [
            ("رابط السيرفر", "https://parental-server-4mms.onrender.com"),
            ("مستودع GitHub", "wh8ps5kwwx-dev/parental-server"),
            ("حزمة الطفل", "com.example.myrana.child"),
            ("حزمة الأم", "com.example.myrana.parent"),
        ],
    )

    # 5.10 Challenges
    heading(doc, "5.10 تحديات التنفيذ وحلولها", 2)
    table(
        doc,
        ["التحدي", "الحل المنفّذ"],
        [
            ("نوم السيرفر على Render", "إيقاظ مسبق عبر المتصفح؛ زيادة مهلة الاتصال"),
            ("استمرار المراقبة بالخلفية", "Foreground Service + BootReceiver"),
            ("فقدان الشبكة", "Room + OutboxRepository"),
            ("تنوع تطبيقات Android", "MonitoredAppRegistry لكل تطبيقات المستخدم"),
            ("الخصوصية مقابل الحماية", "واجهة لعبة + تنبيهات للأم فقط"),
        ],
    )

    # 5.11 Screenshots
    heading(doc, "5.11 لقطات الشاشة (أضيفي من الاختبار الفعلي)", 2)
    shots = [
        "شكل (5-1): شاشة تسجيل الطفل وإرسال CHILD.",
        "شكل (5-2): رسالة Gmail بكود CHILD.",
        "شكل (5-3): شاشة ربط ولي الأمر.",
        "شكل (5-4): لوحة التحكم بعد الربط.",
        "شكل (5-5): لوحة وقت الاستخدام والمؤشرات.",
        "شكل (5-6): شاشة الصلاحيات على جوال الطفل.",
        "شكل (5-7): تنبيه عند البحث عن كلمة خطرة.",
        "شكل (5-8): استجابة السيرفر (status: running).",
    ]
    for s in shots:
        para(doc, f"[{s}]", indent=0)

    heading(doc, "5.12 خلاصة الفصل", 2)
    para(
        doc,
        "تم تنفيذ نظام رقابة أبوية متكامل يجمع بين واجهة طفل مموّهة (أكاديمية "
        "العباقرة) ولوحة أم (MY Rana) وسيرفر Flask سحابي مع ربط Gmail حقيقي. "
        "تغطي طبقة المراقبة وقت الاستخدام، النصوص، المواقع، الملفات، وقت النوم، "
        "والتنبيهات. يمثل هذا الفصل الأساس العملي الذي يدعم الفصول النظرية "
        "والاختبارية في المشروع.",
        indent=0.75,
    )


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = sec.bottom_margin = sec.right_margin = sec.left_margin = Cm(2.5)

    build_chapter5(doc)

    saved = []
    for path in OUT_PATHS:
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            doc.save(path)
            saved.append(path)
        except OSError as e:
            print(f"skip {path}: {e}")

    if not saved:
        fallback = os.path.join(os.path.dirname(__file__), "chapter5_implementation.docx")
        doc.save(fallback)
        saved.append(fallback)

    for p in saved:
        print(f"Saved: {p}")


if __name__ == "__main__":
    main()
